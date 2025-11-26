# rag/document_processor.py
import os
import re
from typing import List, Dict, Any, Optional
from langchain_classic.text_splitter import RecursiveCharacterTextSplitter
from langchain_classic.schema import Document as LangchainDocument
import PyPDF2
import docx
import markdown


class DocumentProcessor:
    """文档处理器，支持多种文件格式的解析和处理"""

    def __init__(self):
        # 初始化文本分割器
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", " ", "", "\."]
        )

    def process_file(self, file_path: str) -> List[LangchainDocument]:
        """处理单个文件，返回文档块列表"""
        file_ext = os.path.splitext(file_path)[1].lower()

        if file_ext == ".pdf":
            text = self._extract_pdf(file_path)
        elif file_ext == ".docx":
            text = self._extract_docx(file_path)
        elif file_ext in [".md", ".markdown"]:
            text = self._extract_markdown(file_path)
        elif file_ext == ".txt":
            text = self._extract_txt(file_path)
        elif file_ext == ".py":
            text = self._extract_python(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")

        # 分割文本
        chunks = self.text_splitter.split_text(text)

        # 创建文档对象
        documents = []
        for i, chunk in enumerate(chunks):
            metadata = {
                "source": file_path,
                "file_name": os.path.basename(file_path),
                "chunk_index": i,
                "total_chunks": len(chunks)
            }
            documents.append(LangchainDocument(page_content=chunk, metadata=metadata))

        return documents

    def _extract_pdf(self, file_path: str) -> str:
        """从PDF文件提取文本"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"PDF提取错误: {e}")
        return text

    def _extract_docx(self, file_path: str) -> str:
        """从DOCX文件提取文本"""
        text = ""
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            print(f"DOCX提取错误: {e}")
        return text

    def _extract_markdown(self, file_path: str) -> str:
        """从Markdown文件提取文本"""
        text = ""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except Exception as e:
            print(f"Markdown提取错误: {e}")
        return text

    def _extract_txt(self, file_path: str) -> str:
        """从TXT文件提取文本"""
        text = ""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as file:
                    text = file.read()
            except Exception as e:
                print(f"TXT提取错误: {e}")
        except Exception as e:
            print(f"TXT提取错误: {e}")
        return text

    def _extract_python(self, file_path: str) -> str:
        """从Python文件提取文本和注释"""
        text = ""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

                # 保留文件头部的注释（通常包含文件说明）
                header_comments = []
                lines = content.split('\n')
                for line in lines:
                    stripped = line.strip()
                    if stripped.startswith('#') or not stripped:
                        header_comments.append(line)
                    else:
                        break

                # 添加文件路径和名称信息
                text = f"# 文件路径: {file_path}\n# 文件名称: {os.path.basename(file_path)}\n\n"
                text += '\n'.join(header_comments) + '\n\n' + content
        except Exception as e:
            print(f"Python文件提取错误: {e}")
        return text

    def process_web_content(self, url: str, content: str, title: str = None) -> List[LangchainDocument]:
        """处理网页内容，返回文档块列表"""
        # 清理网页内容
        # 移除HTML标签
        clean_content = re.sub(r'<[^>]+>', '', content)
        # 移除多余的空白字符
        clean_content = re.sub(r'\s+', ' ', clean_content).strip()

        # 添加URL和标题信息
        if title:
            formatted_content = f"# 标题: {title}\n# URL: {url}\n\n{clean_content}"
        else:
            formatted_content = f"# URL: {url}\n\n{clean_content}"

        # 分割文本
        chunks = self.text_splitter.split_text(formatted_content)

        # 创建文档对象
        documents = []
        for i, chunk in enumerate(chunks):
            metadata = {
                "source": url,
                "title": title,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "type": "web_content"
            }
            documents.append(LangchainDocument(page_content=chunk, metadata=metadata))

        return documents